#coding:utf-8
import requests
import json
import docx

###If you're going to use it a ton, grab a free API key from ocr.space and sub it into the kwargs
###Change the directory to wherever you are saving your screenshots
###Use Snipping Tool on windows, or whatever Apple has for the same purpose to grab screenshots of text.
###You need to grab it block by block. For example, if there are two columns of text beside each other, they will get
###garbled together. You'll see what I mean if you try and use it and it messes up.

###Save your screenshots wherever you want, and then change this variable to that file path:
directory = "C:\Users\Ryan\Desktop\OCR"
###Also, I have it set up so that the names of the screenshots should be 001.jpg, 002.jpg, 003.jpg etc., but
###you can change that below.


def ocr_space_file(filename, overlay=False, api_key='551ce058d588957', language='jpn'):
    """ OCR.space API request with local file.
        Python3.5 - not tested on 2.7
    :param filename: Your file path & name.
    :param overlay: Is OCR.space overlay required in your response.
                    Defaults to False.
    :param api_key: OCR.space API key.
                    Defaults to 'helloworld'.
    :param language: Language code to be used in OCR.
                    List of available language codes can be found on https://ocr.space/OCRAPI
                    Defaults to 'en'.
    :return: Result in JSON format.
    """

    payload = {'isOverlayRequired': overlay,
               'apikey': api_key,
               'language': language,
               }
    with open(filename, 'rb') as f:
        r = requests.post('https://api.ocr.space/parse/image',
                          files={filename: f},
                          data=payload,
                          )
    m = r.content.decode()
    jsonstr = json.loads(m)
    return jsonstr

mydoc = docx.Document()
#change the range/zfill value below based on the number of images
for i in range(1,5):
    try:
        jsonstr = ocr_space_file(filename=f'{directory}\\{str(i).zfill(2)}.jpg', language='jpn')
        mydoc.add_picture(f'{directory}\\{str(i).zfill(2)}.jpg')
        text = jsonstr['ParsedResults'][0]["ParsedText"]
        mydoc.add_paragraph(text.replace("\n",""))
        mydoc.add_page_break()
    except IndexError:
        pass
mydoc.save(f'{directory}\\file.docx')
